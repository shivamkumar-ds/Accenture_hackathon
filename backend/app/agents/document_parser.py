"""
Document text extraction.

Plain functions, one per format, dispatched with a simple if/elif in
extract_text() — no parser factory or plugin registry. There are three
formats to support; a registration abstraction would be complexity with
no present benefit for that.
"""

from dataclasses import dataclass
from pathlib import Path

import pytesseract
from pdf2image import convert_from_path
from pypdf import PdfReader
from docx import Document as DocxDocument

# Below this many characters, a "text-extracted" PDF is treated as if it
# were actually a scanned image (no real text layer) and OCR is used instead.
MIN_NATIVE_TEXT_LENGTH = 50


@dataclass
class ParsedDocument:
    text: str
    used_ocr: bool
    ocr_confidence: float | None  # 0-100 (tesseract's own scale), None when OCR wasn't used


def extract_text(file_path: Path, extension: str) -> ParsedDocument:
    if extension == ".pdf":
        return _extract_pdf(file_path)
    if extension == ".docx":
        return ParsedDocument(text=_extract_docx(file_path), used_ocr=False, ocr_confidence=None)
    if extension in (".png", ".jpg", ".jpeg"):
        return _extract_image(file_path)
    if extension in (".xls", ".xlsx"):
        sheets = extract_spreadsheet_sheets(file_path, extension)
        combined = "\n\n".join(f"=== {name} ===\n{text}" for name, text in sheets)
        return ParsedDocument(text=combined, used_ocr=False, ocr_confidence=None)
    raise ValueError(f"Unsupported file extension for parsing: {extension}")


def extract_spreadsheet_sheets(file_path: Path, extension: str) -> list[tuple[str, str]]:
    """
    Extracts every non-empty sheet in a spreadsheet as (sheet_name, text)
    pairs -- structured/textual, never a binary dump, and never a
    spreadsheet-editing/rendering feature. Each row becomes one line,
    cells joined with " | ", so row/column relationships stay legible to
    an LLM without carrying full binary formatting. Fully blank rows and
    fully blank sheets are dropped rather than passed through as noise.

    .xlsx uses openpyxl (the modern OOXML format). .xls uses xlrd (the
    legacy binary format, which openpyxl cannot read at all -- xlrd is
    the only maintained pure-Python reader for it). Real government
    procurement portals sometimes mislabel a genuinely-OOXML file with a
    .xls extension (observed against a real CPPP tender's bid documents)
    -- if xlrd can't open it as legacy binary, this falls back to
    openpyxl before giving up, rather than failing a file that's
    actually perfectly readable.
    """
    if extension == ".xlsx":
        return [(name, text) for name, text in _extract_xlsx_sheets(file_path) if text.strip()]
    if extension == ".xls":
        try:
            sheets = _extract_xls_sheets(file_path)
        except Exception:
            sheets = _extract_xlsx_sheets(file_path)
        return [(name, text) for name, text in sheets if text.strip()]
    raise ValueError(f"Unsupported spreadsheet extension: {extension}")


def _extract_xlsx_sheets(file_path: Path) -> list[tuple[str, str]]:
    import openpyxl

    workbook = openpyxl.load_workbook(str(file_path), data_only=True, read_only=True)
    try:
        result = []
        for sheet in workbook.worksheets:
            lines = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    lines.append(" | ".join(cells))
            result.append((sheet.title, "\n".join(lines)))
        return result
    finally:
        workbook.close()


def _extract_xls_sheets(file_path: Path) -> list[tuple[str, str]]:
    import xlrd

    workbook = xlrd.open_workbook(str(file_path))
    result = []
    for sheet in workbook.sheets():
        lines = []
        for row_index in range(sheet.nrows):
            row = sheet.row_values(row_index)
            cells = [str(cell).strip() for cell in row if cell not in (None, "") and str(cell).strip()]
            if cells:
                lines.append(" | ".join(cells))
        result.append((sheet.name, "\n".join(lines)))
    return result


def extract_pdf_pages(file_path: Path) -> list[str]:
    """
    Page-indexed extraction for large documents (tenders) where source
    page attribution matters — extract_text() above collapses a whole
    document into one blob, which is fine for a 1-2 page certificate but
    loses page numbers entirely, which M5 needs for provenance.

    No OCR fallback here (unlike extract_text): tenders are realistically
    born-digital procurement documents, not scanned certificates, and
    per-page OCR across a real 300-page tender is a different performance
    conversation than a 2-page certificate. A page with no native text
    returns an empty string rather than being OCR'd.
    """
    reader = PdfReader(str(file_path))
    return [page.extract_text() or "" for page in reader.pages]


def _extract_pdf(file_path: Path) -> ParsedDocument:
    reader = PdfReader(str(file_path))
    native_text = "\n".join(page.extract_text() or "" for page in reader.pages)

    if len(native_text.strip()) >= MIN_NATIVE_TEXT_LENGTH:
        return ParsedDocument(text=native_text, used_ocr=False, ocr_confidence=None)

    # No usable text layer — likely a scanned PDF. Fall back to OCR.
    images = convert_from_path(str(file_path))
    ocr_text_parts = []
    confidences = []
    for image in images:
        # image_to_string preserves line structure — needed for downstream
        # field-matching, which relies on newlines to bound a value. Using
        # image_to_data for text instead (joining its word list) was a real
        # bug: it discards line breaks entirely, caught via testing.
        ocr_text_parts.append(pytesseract.image_to_string(image))
        data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        confidences.extend(c for c in data["conf"] if isinstance(c, (int, float)) and c >= 0)

    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
    return ParsedDocument(text="\n".join(ocr_text_parts), used_ocr=True, ocr_confidence=avg_confidence)


def _extract_docx(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _extract_image(file_path: Path) -> ParsedDocument:
    from PIL import Image

    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    confidences = [c for c in data["conf"] if isinstance(c, (int, float)) and c >= 0]
    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
    return ParsedDocument(text=text, used_ocr=True, ocr_confidence=avg_confidence)
